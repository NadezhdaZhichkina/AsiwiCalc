import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Калькулятор партнёрской прибыли", layout="centered")
st.title("🔢 Калькулятор партнёрской прибыли")

partners = {
    "Партнёр с НДС": {"nds": True},
    "Партнёр без НДС": {"nds": False},
}

partner_name = st.selectbox("Выберите подрядчика:", list(partners.keys()))
partner_nds = partners[partner_name]["nds"]

uploaded_file = st.file_uploader("Загрузите спецификацию (Excel или DOCX):", type=["xlsx", "xls", "docx"])

def parse_file(file):
    if file.name.endswith(".xlsx") or file.name.endswith(".xls"):
        df = pd.read_excel(file)
    elif file.name.endswith(".docx"):
        doc = Document(file)
        rows = []
        for table in doc.tables:
            for row in table.rows:
                cols = [cell.text.strip() for cell in row.cells]
                if len(cols) >= 2:
                    try:
                        cost = float(cols[1].replace(",", ".").replace(" ", ""))
                        rows.append([cols[0], cost])
                    except:
                        continue
        df = pd.DataFrame(rows, columns=["Наименование", "Стоимость"])
    else:
        df = pd.DataFrame()
    return df

def find_price_column(df):
    for col in df.columns:
        if isinstance(col, str) and any(x in col.lower() for x in ["стоим", "цена"]):
            return col
    return None

def generate_docx(table_df, total):
    doc = Document()
    doc.add_heading('Спецификация для клиента', 0)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Услуга'
    hdr_cells[1].text = 'Цена с НДС (₽)'
    for i, row in table_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Услуга'])
        cells[1].text = f"{row['Цена с НДС']:,.2f} ₽".replace(",", " ").replace(".", ",")
    doc.add_paragraph()
    doc.add_paragraph(f"{'Итого с НДС: ' + f'{total:,.2f} ₽'.replace(',', ' ').replace('.', ',')}")
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

if uploaded_file:
    df_spec = parse_file(uploaded_file)
    if df_spec.empty:
        st.error("❌ Не удалось распознать таблицу.")
    else:
        st.subheader("📄 Считанная спецификация:")
        price_col = find_price_column(df_spec)
        if not price_col:
            st.error("Не найдена колонка с ценой.")
        else:
            formatted_df = df_spec.copy()
            formatted_df[price_col] = formatted_df[price_col].apply(lambda x: f"{x:,.2f}".replace(",", " ").replace(".", ","))
            st.dataframe(formatted_df)

            total_partner_sum = df_spec[price_col].sum()
            first_col = df_spec.columns[0]
            st.markdown(f"**Сумма от подрядчика:** `{f'{total_partner_sum:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")

            st.markdown("### ➕ Расчёт чистой прибыли")

            if "client_sum" not in st.session_state:
                st.session_state.client_sum = 0.0
            if "desired_profit" not in st.session_state:
                st.session_state.desired_profit = 0.0

            def on_client_sum_change():
                st.session_state.desired_profit = 0.0

            st.number_input(
                "Сумма, которую планируем выставить клиенту (включает НДС):",
                step=1000.0,
                key="client_sum",
                on_change=on_client_sum_change
            )

            st.number_input(
                "Желаемая чистая прибыль (обратный расчёт):",
                step=1000.0,
                key="desired_profit"
            )

            client_sum = st.session_state.client_sum
            desired_profit = st.session_state.desired_profit

            nds_included = partner_nds and st.checkbox("Сумма в спецификации включает НДС", value=True)

            col1, col2 = st.columns(2)
            with col1:
                calc_button = st.button("🔁 Пересчитать прибыль")
            with col2:
                spec_button = st.button("📋 Показать спецификацию и выгрузить в DOCX/Excel")

            if desired_profit > 0:
                d = desired_profit
                if partner_nds:
                    if nds_included:
                        nds_sub = total_partner_sum * 20 / 120
                        net_sub = total_partner_sum - nds_sub
                    else:
                        net_sub = total_partner_sum
                    nds_loss = net_sub * 0.2 * 0.75
                    tax_base = d / 0.95
                    net_client = net_sub + nds_loss + tax_base
                    x = net_client * 1.2
                else:
                    net_sum = d / 0.95 + total_partner_sum
                    x = net_sum * 1.2
                client_sum = x
                st.info("🧾 Чтобы получить **{} ₽** прибыли, нужно выставить клиенту: **{} ₽**".format(
                    f"{d:,.2f}".replace(",", " ").replace(".", ","),
                    f"{x:,.2f}".replace(",", " ").replace(".", ","))
                )

            if calc_button and client_sum > 0:
                st.markdown("### 📊 Расчёт прибыли")
                client_nds = client_sum * 20 / 120
                client_net = client_sum - client_nds

                if partner_nds:
                    if nds_included:
                        nds_sub = total_partner_sum * 20 / 120
                        net_sub = total_partner_sum - nds_sub
                    else:
                        net_sub = total_partner_sum
                        nds_sub = net_sub * 0.20
                        total_partner_sum = net_sub + nds_sub

                    nds_loss = nds_sub * 0.75
                    direct_costs = net_sub + nds_loss
                    tax_base = client_net - direct_costs
                    tax = tax_base * 0.05
                    profit = tax_base - tax

                    st.markdown(f"- **НДС подрядчика:** `{f'{nds_sub:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("20% от стоимости, если подрядчик работает с НДС. Учитывается для расчёта убытка по НДС.")

                    st.markdown(f"- **Нетто подрядчику:** `{f'{net_sub:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Сумма, которую реально получает подрядчик за вычетом НДС.")

                    st.markdown(f"- **НДС клиента:** `{f'{client_nds:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("20% от выставленной клиенту суммы (если она включает НДС).")

                    st.markdown(f"- **Нетто от клиента:** `{f'{client_net:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Сумма, полученная от клиента без НДС.")

                    st.markdown(f"- **Убыток по НДС (75%):** `{f'{nds_loss:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Часть входящего НДС, которую нельзя зачесть — условный убыток. Берётся 75% от НДС подрядчика.")

                    st.markdown(f"- **Прямые расходы:** `{f'{direct_costs:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Нетто подрядчику + убыток по НДС. Всё, что компания тратит напрямую.")

                    st.markdown(f"- **Налоговая база:** `{f'{tax_base:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Разница между доходом от клиента и прямыми расходами.")

                    st.markdown(f"- **Налог на прибыль (5%):** `{f'{tax:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("5% от налоговой базы — упрощённый расчёт налога на прибыль.")

                    st.success(f"💰 **Чистая прибыль:** `{f'{profit:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("То, что остаётся после всех налогов и расходов.")
                else:
                    tax_base = client_net - total_partner_sum
                    tax = tax_base * 0.05
                    profit = tax_base - tax

                    st.markdown(f"- **Нетто подрядчику (без НДС):** `{f'{total_partner_sum:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.caption("Полная сумма, уплаченная подрядчику (без НДС).")

                    st.markdown(f"- **НДС клиента:** `{f'{client_nds:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.markdown(f"- **Нетто от клиента:** `{f'{client_net:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.markdown(f"- **Налоговая база:** `{f'{tax_base:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.markdown(f"- **Налог на прибыль (5%):** `{f'{tax:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")
                    st.success(f"💰 **Чистая прибыль:** `{f'{profit:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`")

            if spec_button:
                st.markdown("### 📑 Спецификация для клиента")
                spec_df = df_spec.copy()
                total_original = spec_df[price_col].sum()

                if client_sum <= 0 or total_original == 0:
                    st.warning("Введите сумму для клиента и убедитесь, что спецификация не пустая.")
                else:
                    k = client_sum / total_original
                    spec_df["Цена с НДС"] = spec_df[price_col] * k
                    spec_df["Цена с НДС"] = spec_df["Цена с НДС"].round(2)
                    spec_display = spec_df[[first_col, "Цена с НДС"]].rename(columns={first_col: "Услуга"})

                    formatted_display = spec_display.copy()
                    formatted_display["Цена с НДС"] = formatted_display["Цена с НДС"].apply(lambda x: f"{x:,.2f}".replace(",", " ").replace(".", ","))

                    st.dataframe(formatted_display)
                    total_for_client = spec_display["Цена с НДС"].sum()
                    st.markdown(f"💼 **Итоговая сумма для клиента (с НДС): `{f'{total_for_client:,.2f} ₽'.replace(',', ' ').replace('.', ',')}`**")

                    docx_file = generate_docx(spec_display, total_for_client)
                    st.download_button("💾 Скачать спецификацию (DOCX)", docx_file, file_name="Спецификация_для_клиента.docx")

                    excel_io = BytesIO()
                    with pd.ExcelWriter(excel_io, engine="openpyxl") as writer:
                        spec_display.to_excel(writer, index=False, sheet_name="Спецификация")
                        worksheet = writer.sheets["Спецификация"]
                        worksheet.cell(row=len(spec_display) + 2, column=1, value="Итого:")
                        worksheet.cell(row=len(spec_display) + 2, column=2, value=total_for_client)
                    excel_io.seek(0)
                    st.download_button("📥 Скачать спецификацию (Excel)", excel_io, file_name="Спецификация_для_клиента.xlsx")
